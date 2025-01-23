from docx import Document

def create_word_document(content, output_path):
    try:
        doc = Document()
        doc.add_heading('Scraped Content', level=1)

        for item in content:
            doc.add_paragraph(item)

        doc.save(output_path)
        print(f"Document saved to {output_path}")
    except Exception as e:
        print(f"Error creating Word document: {e}")

if __name__ == "__main__":
    sample_content = ["Title: Example", "Body: This is an example."]
    create_word_document(sample_content, "output.docx")
